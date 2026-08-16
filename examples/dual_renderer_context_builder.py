"""
Patrón: un único "context builder" arma un diccionario con TODOS los datos
ya procesados del informe, y dos renderers independientes (HTML->PDF y
Word) lo consumen sin volver a tocar la lógica de negocio.

Extracto representativo (renombrado y simplificado) - en el sistema real
cada informe se genera en dos formatos, PDF (WeasyPrint) y Word
(python-docx), y ambos tienen que ser fieles el uno al otro: mismas
secciones, mismos datos, mismo orden. La tentación es escribir la lógica
de "qué va en el informe" dos veces, una por formato - eso garantiza que
con el tiempo diverjan (alguien corrige un caso borde en el PDF y se
olvida del Word). La solución real: una sola función construye el
contexto (qué secciones van, con qué datos, ya validados), y cada
renderer solo sabe *dibujar*, no *decidir*.

Esto también obliga a resolver un problema real de Word: python-docx no
tiene una forma "linda" de controlar ciertos detalles de formato (márgenes
internos de celda, por ejemplo) - hay que bajar al XML de OOXML a mano
con `OxmlElement`. Se muestra ese fragmento real abajo, no es un adorno.
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# --- 1. Contexto único, construido una sola vez ---------------------------

def construir_contexto(cliente: dict, mediciones: list[dict]) -> dict:
    """Un solo lugar decide QUÉ va en el informe y con qué forma. Los
    renderers de abajo no vuelven a tomar ninguna decisión de negocio,
    solo dibujan lo que ya viene resuelto acá."""

    mediciones_ok = [m for m in mediciones if m.get("valor") is not None]

    return {
        "titulo": f"Informe de mantención - {cliente['nombre']}",
        "cliente": cliente["nombre"],
        "filas": [
            {"equipo": m["equipo"], "valor": m["valor"], "unidad": m["unidad"]}
            for m in mediciones_ok
        ],
        "hay_observaciones": any(m.get("observacion") for m in mediciones_ok),
    }


# --- 2. Renderer A: HTML (en el sistema real, WeasyPrint lo convierte a PDF)

def renderizar_html(contexto: dict) -> str:
    filas_html = "".join(
        f"<tr><td>{f['equipo']}</td><td>{f['valor']} {f['unidad']}</td></tr>"
        for f in contexto["filas"]
    )
    return (
        f"<h1>{contexto['titulo']}</h1>"
        f"<table><tbody>{filas_html}</tbody></table>"
    )


# --- 3. Renderer B: Word, con una celda de margen interno ajustado a mano -

def _fijar_margenes_celda(celda, top=40, bottom=40, left=80, right=80):
    """python-docx no expone `cell.margins` - hay que inyectar el
    elemento OOXML <w:tcMar> directamente en las propiedades de la celda.
    Esto es real, no un ejemplo de juguete: sin esto las tablas del
    informe salían con el texto pegado al borde en Word (se veía bien en
    la vista previa de LibreOffice pero mal en Word real, que sí respeta
    estos márgenes explícitos en vez de usar un default razonable)."""

    tc_pr = celda._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")

    for lado, valor in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        nodo = OxmlElement(f"w:{lado}")
        nodo.set(qn("w:w"), str(valor))
        nodo.set(qn("w:type"), "dxa")
        tc_mar.append(nodo)

    tc_pr.append(tc_mar)


def renderizar_word(contexto: dict) -> Document:
    doc = Document()
    doc.add_heading(contexto["titulo"], level=1)

    tabla = doc.add_table(rows=1, cols=2)
    encabezado = tabla.rows[0].cells
    encabezado[0].text, encabezado[1].text = "Equipo", "Valor"

    for fila in contexto["filas"]:
        celdas = tabla.add_row().cells
        celdas[0].text = fila["equipo"]
        celdas[1].text = f"{fila['valor']} {fila['unidad']}"
        for celda in celdas:
            _fijar_margenes_celda(celda)

    return doc


if __name__ == "__main__":
    cliente = {"nombre": "Edificio Ejemplo"}
    mediciones = [
        {"equipo": "Chiller 1", "valor": 6.2, "unidad": "bar", "observacion": None},
        {"equipo": "Torre 1", "valor": None, "unidad": "°C", "observacion": None},  # se descarta
        {"equipo": "Torre 2", "valor": 28.4, "unidad": "°C", "observacion": "ok"},
    ]

    contexto = construir_contexto(cliente, mediciones)
    print("--- HTML ---")
    print(renderizar_html(contexto))

    doc = renderizar_word(contexto)
    doc.save("informe_ejemplo.docx")
    print("--- Word ---")
    print("informe_ejemplo.docx generado con", len(contexto["filas"]), "filas")
